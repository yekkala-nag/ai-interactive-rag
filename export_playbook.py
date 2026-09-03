#!/usr/bin/env python3
"""
export_playbook.py
Enterprise AI Playbook Word (.docx) Generator
Curated by: Nagaraj Y
Disclaimer: This is only for Education purpose

Generates a formatted Microsoft Word (.docx) document
covering Enterprise AI Architecture, Token Governance, Agentic Patterns,
Distributed Computing, Security Guardrails, and Stack Adaptations.
"""

import os
import sys

def create_enterprise_playbook_docx(output_filename="Enterprise_AI_Playbook.docx"):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("Required package 'python-docx' is not installed.")
        print("Run: pip install python-docx")
        return False

    doc = Document()

    # --- Title Page ---
    doc.add_paragraph('\n\n\n')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Enterprise AI Playbook:\nToken Governance, Agentic Patterns,\nObservability & Distributed Computing')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('\n')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Production Architecture Blueprint for Engineering & Technology Leaders')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('\n')
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run('Curated by: Nagaraj Y\nDisclaimer: ⚠️ This is only for Education purpose.')
    author_run.font.size = Pt(11)
    author_run.font.italic = True
    author_run.font.color.rgb = RGBColor(180, 100, 20)

    doc.add_page_break()

    # --- Helper Styling Functions ---
    def add_heading_styled(text, level=1):
        heading = doc.add_heading(text, level=level)
        for r in heading.runs:
            r.font.color.rgb = RGBColor(0, 51, 102)
        return heading

    def add_code_block(code_text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 30, 30)

    # --- Table of Contents Overview ---
    add_heading_styled('Executive Summary & Architectural Pillars', level=1)
    doc.add_paragraph(
        'This playbook provides an end-to-end blueprint for engineering and enterprise architecture teams '
        'deploying large language models, autonomous agentic systems, and distributed RAG pipelines at scale. '
        'It synthesizes strategies across five core operational domains:'
    )

    bullet_points = [
        ("Production Deployment Architecture:", "Containerized microservices with FastAPI, Redis Stack, and Qdrant deployed on Kubernetes with HPA and Secrets Vault."),
        ("Token Governance & Observability:", "LangSmith run-tree tracing, Prometheus metrics scraping, and OpenTelemetry spans quantifying ROI and cache hits."),
        ("Agentic System Patterns:", "ReAct loops, Plan-and-Execute StateGraphs, Multi-Agent reviewer cycles, and Supervisor-Worker hierarchical delegation."),
        ("Distributed Workload Computing:", "Celery chord Map-Reduce tasks, Ray distributed actor clusters, and Kafka event streaming."),
        ("Security & Guardrails:", "Presidio PII redaction, LLM Guard prompt injection defense, and RBAC pre-filtered vector database retrieval."),
        ("Enterprise Tech Stack Adaptations:", "Java/Spring AI implementations and AWS Bedrock cloud-native agents with Step Functions.")
    ]

    for title_text, desc in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(title_text + ' ')
        r_bold.bold = True
        p.add_run(desc)

    doc.add_paragraph('\n')

    # --- Part 1: Deployment ---
    add_heading_styled('Part 1: Containerized Microservices Deployment (Docker & K8s)', level=1)
    doc.add_paragraph(
        'Running production RAG requires decoupling the orchestration runtime from caching and vector persistence. '
        'Below is the standard docker-compose.yml configuration for local development and staging:'
    )
    add_code_block(
        "version: '3.8'\nservices:\n  ai-api:\n    build: .\n    ports: [\"8000:8000\"]\n"
        "    environment:\n      - OPENAI_API_KEY=${OPENAI_API_KEY}\n      - REDIS_HOST=redis\n"
        "      - QDRANT_HOST=qdrant\n    depends_on: [redis, qdrant]\n\n"
        "  redis:\n    image: redis/redis-stack:latest\n    ports: [\"6379:6379\", \"8001:8001\"]\n\n"
        "  qdrant:\n    image: qdrant/qdrant:latest\n    ports: [\"6333:6333\", \"6334:6334\"]"
    )

    # --- Part 2: Security & Guardrails ---
    add_heading_styled('Part 2: Security, Compliance & Guardrails', level=1)
    doc.add_paragraph('Enterprise deployments mandate pre-LLM PII redaction and active prompt injection defense.')
    
    add_heading_styled('2.1 Microsoft Presidio PII Redaction', level=2)
    add_code_block(
        "from presidio_analyzer import AnalyzerEngine\n"
        "from presidio_anonymizer import AnonymizerEngine\n\n"
        "analyzer = AnalyzerEngine()\n"
        "anonymizer = AnonymizerEngine()\n\n"
        "def redact_pii(text: str) -> str:\n"
        "    results = analyzer.analyze(text=text, language='en')\n"
        "    return anonymizer.anonymize(text=text, analyzer_results=results).text"
    )

    add_heading_styled('2.2 Prompt Injection Defense (LLM Guard)', level=2)
    add_code_block(
        "from llm_guard.input_scanners import PromptInjection, Toxicity\n"
        "from llm_guard import scan_prompt\n\n"
        "scanners = [PromptInjection(threshold=0.5), Toxicity()]\n"
        "def validate_input(user_input: str):\n"
        "    sanitized, valid, score = scan_prompt(scanners, user_input)\n"
        "    if not all(valid.values()):\n"
        "        raise ValueError(f'Security violation: {score}')\n"
        "    return sanitized"
    )

    add_heading_styled('2.3 Role-Based Access Control (RBAC) in Vector Search', level=2)
    add_code_block(
        "def search_with_rbac(user_id: str, query: str):\n"
        "    authorized_tags = get_user_permissions(user_id)\n"
        "    return vectorstore.similarity_search(\n"
        "        query=query, k=5, filter={'access_level': {'$in': authorized_tags}}\n"
        "    )"
    )

    # --- Part 3: Stack Adaptations ---
    add_heading_styled('Part 3: Tech Stack Adaptations (Spring AI & AWS Bedrock)', level=1)
    doc.add_paragraph('For enterprise ecosystems utilizing Java or AWS cloud-native primitives:')

    add_heading_styled('3.1 Java & Spring AI RAG Implementation', level=2)
    add_code_block(
        "@Service\npublic class EnterpriseRagService {\n"
        "    @Autowired private ChatClient chatClient;\n"
        "    @Autowired private VectorStore vectorStore;\n\n"
        "    public String queryWithRag(String question, String userId) {\n"
        "        List<String> roles = getUserRoles(userId);\n"
        "        SearchRequest searchRequest = SearchRequest.builder()\n"
        "                .query(question).filterExpression(\"roles in \" + roles).topK(5).build();\n"
        "        List<Document> docs = vectorStore.similaritySearch(searchRequest);\n"
        "        return chatClient.prompt(new Prompt(question)).call().content();\n"
        "    }\n}"
    )

    add_heading_styled('3.2 AWS Bedrock Streaming Agent Invocation', level=2)
    add_code_block(
        "import boto3\n"
        "bedrock_agent_runtime = boto3.client('bedrock-agent-runtime')\n"
        "response = bedrock_agent_runtime.invoke_agent(\n"
        "    agentId='YOUR_AGENT_ID', agentAliasId='YOUR_ALIAS_ID',\n"
        "    sessionId='session-123', inputText='Reset my password'\n"
        ")"
    )

    # --- Part 4: Implementation Roadmap ---
    add_heading_styled('Part 4: 10-Week Enterprise Implementation Roadmap', level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['Timeframe', 'Phase Focus', 'Key Enterprise Deliverables']
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text

    roadmap_data = [
        ('Weeks 1-2', 'Audit & Core Observability', 'Presidio PII filters, LangSmith tracing, Prometheus endpoints'),
        ('Weeks 3-4', 'Enterprise Agent Patterns', 'LangGraph StateGraph, ReAct loops, typed @tool definitions'),
        ('Weeks 5-6', 'Distributed Processing', 'Celery chord Map-Reduce, Ray actor pools, Redis caching'),
        ('Weeks 7-8', 'Streaming & Message Queues', 'Kafka query/processing/response topics, dead-letter queues'),
        ('Weeks 9-10', 'Autoscaling & Hardening', 'Consistent Hash Router, K8s HPA custom metrics, RBAC audit')
    ]

    for row_idx, row_data in enumerate(roadmap_data, start=1):
        for col_idx, cell_value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_value

    doc.save(output_filename)
    print(f"Successfully generated formatted Word document: {output_filename}")
    return True

if __name__ == '__main__':
    filename = sys.argv[1] if len(sys.argv) > 1 else "Enterprise_AI_Playbook.docx"
    create_enterprise_playbook_docx(filename)
